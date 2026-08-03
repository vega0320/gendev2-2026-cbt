from __future__ import annotations

import hashlib
import json
import os
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT.parent
SOURCE_DOCX = next(
    path for path in SOURCE_DIR.glob("*.docx") if not path.name.startswith("~$")
)
SOURCE_PDF = ROOT / "source_render" / "source.pdf"
OUTPUT_DATA = ROOT / "data" / "questions.json"
OUTPUT_REPORT = ROOT / "work" / "extraction_report.json"
ASSET_DIR = ROOT / "site" / "assets" / "questions"
FONT_PATH = Path(r"C:\Windows\Fonts\malgun.ttf")

YEAR_RE = re.compile(r"^(2025|2023|2022|2021|2020)(?:\s.*)?$")
LECTURE_RE = re.compile(r"^\d{1,2}(?:-\d)?$")
QUESTION_RE = re.compile(r"^(\d{1,3})([\.,])\s*(.*)$", re.S)
ANSWER_RE = re.compile(r"^(\d{1,3})\.\s*([1-5?])(?:\s*,\s*([1-5]))?(.*)$")
CHOICE_RE = re.compile(r"^(?:[①②③④⑤]|[1-5][\).])\s*(.*)$", re.S)


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def paragraph_text(element) -> str:
    pieces: list[str] = []
    for node in element.iter():
        local = node.tag.rsplit("}", 1)[-1]
        if local == "t":
            pieces.append(node.text or "")
        elif local in {"br", "cr"}:
            pieces.append("\n")
        elif local == "tab":
            pieces.append("\t")
    lines = [clean_text(line) for line in "".join(pieces).splitlines()]
    return "\n".join(line for line in lines if line).strip()


def year_value(text: str) -> str | None:
    match = YEAR_RE.fullmatch(clean_text(text))
    return match.group(1) if match else None


def table_rows(element) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in element.findall(qn("w:tr")):
        cells: list[str] = []
        for cell in row.findall(qn("w:tc")):
            pieces = [paragraph_text(p) for p in cell.findall(qn("w:p"))]
            cells.append(clean_text(" ".join(piece for piece in pieces if piece)))
        rows.append(cells)
    return rows


def image_blips(element) -> list[str]:
    result: list[str] = []
    for blip in element.iter(qn("a:blip")):
        relationship_id = blip.get(qn("r:embed"))
        if relationship_id:
            result.append(relationship_id)
    return result


def make_blocks(document: Document) -> list[dict]:
    blocks: list[dict] = []
    for body_index, element in enumerate(document.element.body.iterchildren()):
        kind = element.tag.rsplit("}", 1)[-1]
        if kind == "p":
            blocks.append(
                {
                    "bodyIndex": body_index,
                    "type": "paragraph",
                    "text": paragraph_text(element),
                    "imageRels": image_blips(element),
                }
            )
        elif kind == "tbl":
            blocks.append(
                {
                    "bodyIndex": body_index,
                    "type": "table",
                    "rows": table_rows(element),
                    "imageRels": image_blips(element),
                }
            )
    return blocks


def lecture_header(block: dict) -> tuple[str, str] | None:
    if block["type"] != "table" or len(block["rows"]) != 1:
        return None
    row = block["rows"][0]
    if len(row) != 2 or not LECTURE_RE.fullmatch(row[0]):
        return None
    return row[0], row[1]


def professor_table(blocks: list[dict]) -> tuple[list[str], dict[str, dict]]:
    table = next(block for block in blocks if block["type"] == "table" and len(block["rows"]) == 42)
    header = table["rows"][0]
    years = header[2:]
    lectures: dict[str, dict] = {}
    for row in table["rows"][1:]:
        number = row[0].zfill(2)
        professors = {year: row[index + 2] for index, year in enumerate(years)}
        lectures[number] = {
            "number": number,
            "title": row[1],
            "professors": professors,
        }
    return years, lectures


def header_positions(blocks: list[dict]) -> list[tuple[int, str, str]]:
    result = []
    for index, block in enumerate(blocks):
        header = lecture_header(block)
        if header:
            result.append((index, header[0].zfill(2), header[1]))
    return result


def parse_answers(blocks: list[dict], positions: list[tuple[int, str, str]]) -> tuple[int, dict]:
    first_repeat = next(
        index
        for index in range(1, len(positions))
        if positions[index][1] == "01"
    )
    answer_positions = positions[first_repeat:]
    answer_start = answer_positions[0][0]
    answers: dict[str, dict[str, dict[str, dict]]] = defaultdict(lambda: defaultdict(dict))
    for pos_index, (start, lecture_number, _) in enumerate(answer_positions):
        stop = answer_positions[pos_index + 1][0] if pos_index + 1 < len(answer_positions) else len(blocks)
        year = None
        for block in blocks[start + 1 : stop]:
            if block["type"] != "paragraph":
                continue
            text = block["text"]
            parsed_year = year_value(text)
            if parsed_year:
                year = parsed_year
                continue
            if not year:
                continue
            match = ANSWER_RE.fullmatch(text)
            if not match:
                continue
            number, first, second, note = match.groups()
            values = [] if first == "?" else [int(first)]
            if second:
                values.append(int(second))
            status = "족보 정답"
            if first == "?":
                status = "정답 불명"
            elif second or "복수" in note:
                status = "족보 복수정답"
            if "수정" in note:
                status = "족보 수정 표시"
            answers[lecture_number][year][str(int(number))] = {
                "answers": values,
                "status": status,
                "raw": text,
                "note": clean_text(note.strip(" ,()")),
            }
    return answer_start, answers


def related_image(document: Document, relationship_id: str) -> tuple[bytes, str]:
    relationship = document.part.rels[relationship_id]
    part = relationship.target_part
    extension = Path(str(part.partname)).suffix.lower() or ".png"
    return part.blob, extension


def normalized(value: str) -> str:
    return re.sub(r"\s+", "", value or "").lower()


def pdf_pages() -> list[str]:
    reader = PdfReader(str(SOURCE_PDF))
    return [normalized(page.extract_text() or "") for page in reader.pages]


def find_source_pages(question: dict, pages: list[str], hint: int) -> list[int]:
    needle = normalized(question["stem"][:110])
    if not needle:
        return [hint]
    short = needle[:50]
    direct = [index + 1 for index, text in enumerate(pages) if short in text]
    if direct:
        return direct[:2]
    tokens = [normalized(token) for token in re.findall(r"[가-힣A-Za-z0-9]{4,}", question["stem"])[:12]]
    scores = []
    for index, text in enumerate(pages):
        score = sum(1 for token in tokens if token and token in text)
        scores.append((score, -abs((index + 1) - hint), index + 1))
    best = max(scores) if scores else (0, 0, hint)
    return [best[2]] if best[0] >= 2 else [hint]


def split_question_entries(entries: list[dict]) -> tuple[str, list[str], list[dict], str]:
    expanded: list[dict] = []
    for item in entries:
        if item["type"] != "text" or "\n" not in item.get("text", ""):
            expanded.append(item)
            continue
        lines = [clean_text(line) for line in item["text"].splitlines() if clean_text(line)]
        for line_index, line in enumerate(lines):
            expanded.append(
                {
                    "type": "text",
                    "text": line,
                    "imageRels": list(item.get("imageRels", [])) if line_index == 0 else [],
                }
            )
    entries = expanded
    paragraph_positions = [index for index, item in enumerate(entries) if item["type"] == "text" and item["text"]]
    if not paragraph_positions:
        return "", [], entries, "선지 없음"
    explicit = []
    for index in paragraph_positions:
        match = CHOICE_RE.match(entries[index]["text"])
        if match:
            explicit.append((index, clean_text(match.group(1))))
    if len(explicit) >= 5:
        selected = explicit[-5:]
        choice_indexes = {index for index, _ in selected}
        choices = [text for _, text in selected]
        status = "선지 표지 분리"
    elif len(paragraph_positions) >= 6:
        choice_indexes = set(paragraph_positions[-5:])
        choices = [entries[index]["text"] for index in paragraph_positions[-5:]]
        status = "마지막 5문단 선지 분리—원본 대조 필요"
    else:
        choice_indexes = set()
        choices = []
        status = "선지 분리 실패"
    content = [item for index, item in enumerate(entries) if index not in choice_indexes]
    text_parts = [item["text"] for item in content if item["type"] == "text" and item["text"]]
    stem = clean_text(" ".join(text_parts))
    if not choices:
        inline = re.match(
            r"^(.*?)(?:①|1\))\s*(.*?)(?:②|2\))\s*(.*?)(?:③|3\))\s*(.*?)(?:④|4\))\s*(.*?)(?:⑤|5\))\s*(.*)$",
            stem,
            re.S,
        )
        if inline:
            stem = clean_text(inline.group(1))
            choices = [clean_text(inline.group(index)) for index in range(2, 7)]
            content = [item for item in content if item["type"] != "text"]
            if stem:
                content.insert(0, {"type": "text", "text": stem, "imageRels": []})
            status = "한 문단 안 선지 수동 규칙 분리"
    return stem, choices, content, status


def write_question_images(document: Document, question_id: str, entries: list[dict]) -> list[str]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets: list[str] = []
    seen: set[str] = set()
    counter = 0
    for entry in entries:
        for relationship_id in entry.pop("imageRels", []):
            blob, extension = related_image(document, relationship_id)
            digest = hashlib.sha256(blob).hexdigest()
            if digest in seen:
                continue
            seen.add(digest)
            counter += 1
            filename = f"{question_id}-{counter:02d}{extension}"
            path = ASSET_DIR / filename
            path.write_bytes(blob)
            assets.append(f"assets/questions/{filename}")
    return assets


def wrap_cell(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    for source_line in str(text).splitlines() or [""]:
        current = ""
        for character in source_line:
            candidate = current + character
            if draw.textlength(candidate, font=font) <= width or not current:
                current = candidate
            else:
                lines.append(current)
                current = character
        lines.append(current)
    return lines or [""]


def render_table_png(question_id: str, table_index: int, rows: list[list[str]]) -> str:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{question_id}-table-{table_index:02d}.png"
    output = ASSET_DIR / filename
    if not rows:
        Image.new("RGB", (600, 80), "white").save(output)
        return f"assets/questions/{filename}"
    columns = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (columns - len(row)) for row in rows]
    font = ImageFont.truetype(str(FONT_PATH), 28)
    scratch = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(scratch)
    total_width = 1500
    natural = []
    for column in range(columns):
        longest = max((draw.textlength(row[column], font=font) for row in normalized_rows), default=100)
        natural.append(max(140, min(620, int(longest) + 44)))
    scale = total_width / sum(natural)
    widths = [max(120, int(width * scale)) for width in natural]
    widths[-1] += total_width - sum(widths)
    padding_x = 18
    padding_y = 14
    line_height = 40
    wrapped_rows = []
    heights = []
    for row in normalized_rows:
        wrapped = [wrap_cell(draw, cell, font, widths[index] - padding_x * 2) for index, cell in enumerate(row)]
        wrapped_rows.append(wrapped)
        heights.append(max(64, max(len(lines) for lines in wrapped) * line_height + padding_y * 2))
    image = Image.new("RGB", (total_width + 2, sum(heights) + 2), "white")
    draw = ImageDraw.Draw(image)
    y = 1
    for row_index, wrapped in enumerate(wrapped_rows):
        x = 1
        if row_index == 0 and len(rows) > 1:
            draw.rectangle((1, y, total_width + 1, y + heights[row_index]), fill="#f3f4f6")
        for column, lines in enumerate(wrapped):
            width = widths[column]
            draw.rectangle((x, y, x + width, y + heights[row_index]), outline="#111827", width=2)
            text_height = len(lines) * line_height
            text_y = y + max(padding_y, (heights[row_index] - text_height) // 2)
            for line in lines:
                draw.text((x + padding_x, text_y), line, font=font, fill="#111827")
                text_y += line_height
            x += width
        y += heights[row_index]
    image.save(output, optimize=True)
    return f"assets/questions/{filename}"


def convert_tables_to_images(question_id: str, content: list[dict]) -> tuple[list[dict], list[str]]:
    converted = []
    assets = []
    table_index = 0
    for item in content:
        if item["type"] != "table":
            converted.append(item)
            continue
        table_index += 1
        src = render_table_png(question_id, table_index, item["rows"])
        converted.append({"type": "image", "src": src, "alt": f"문항 표 {table_index}"})
        assets.append(src)
    return converted, assets


def importance(lecture: dict, year: str) -> str:
    current = clean_text(lecture["professors"].get("2026", ""))
    previous = clean_text(lecture["professors"].get(year, ""))
    if current and previous and current == previous:
        return "high"
    if current and previous and (current in previous or previous in current):
        return "high"
    return "standard"


def parse_questions(document: Document, blocks: list[dict], lectures: dict, positions, answer_start, answers):
    question_positions = [position for position in positions if position[0] < answer_start]
    questions = []
    unresolved = []
    page_texts = pdf_pages()

    for pos_index, (start, lecture_number, source_title) in enumerate(question_positions):
        stop = question_positions[pos_index + 1][0] if pos_index + 1 < len(question_positions) else answer_start
        professor_key = lecture_number.split("-", 1)[0]
        lecture = lectures[professor_key]
        year = None
        active = None

        def finish_active():
            nonlocal active
            if not active:
                return
            stem, choices, content, choice_status = split_question_entries(active["entries"])
            base_id = f"gendev2-{lecture_number}-{active['year']}-q{active['number']:03d}"
            duplicate_index = 1 + sum(question["id"].startswith(base_id) for question in questions)
            question_id = base_id if duplicate_index == 1 else f"{base_id}-v{duplicate_index}"
            active["id"] = question_id
            active["stem"] = stem
            active["choices"] = choices
            raw_content = [
                {key: value for key, value in item.items() if key != "imageRels"}
                for item in content
                if item.get("type") != "text" or item.get("text")
            ]
            active["content"], table_assets = convert_tables_to_images(question_id, raw_content)
            active["choiceStatus"] = choice_status
            active["assets"] = write_question_images(document, question_id, active["entries"]) + table_assets
            if not active["choices"] or not any(clean_text(choice).replace("(불확실)", "") for choice in active["choices"]):
                active["choices"] = [f"그림 속 {number}번 선지" for number in range(1, 6)]
                active["choiceStatus"] = "선지 내용이 그림에 포함됨"
            elif any(not clean_text(choice) for choice in active["choices"]):
                active["choices"] = [
                    choice if clean_text(choice) else f"그림 속 {number}번 선지"
                    for number, choice in enumerate(active["choices"], 1)
                ]
                active["choiceStatus"] = "일부 선지 내용이 그림에 포함됨"
            elif len(active["choices"]) == 4:
                active["choices"].append("원본에 ⑤ 선지가 없음")
                active["choiceStatus"] = "원본에 4개 선지만 존재—⑤ 검수 필요"
            if duplicate_index > 1:
                active["duplicateStatus"] = f"같은 강의·연도·번호의 두 번째 복원 문제(v{duplicate_index})"
            active.pop("entries", None)
            active["sourcePages"] = find_source_pages(active, page_texts, active["pageHint"])
            active.pop("pageHint", None)
            questions.append(active)
            active = None

        for block in blocks[start + 1 : stop]:
            parsed_year = year_value(block["text"]) if block["type"] == "paragraph" else None
            if parsed_year:
                finish_active()
                year = parsed_year
                continue
            if not year:
                continue
            expected = answers.get(lecture_number, {}).get(year, {})
            match = QUESTION_RE.match(block["text"]) if block["type"] == "paragraph" else None
            if match and match.group(2) == "," and len(clean_text(match.group(3))) < 30:
                match = None
            matched_year = year
            if match and str(int(match.group(1))) not in expected:
                candidate_years = [
                    candidate_year
                    for candidate_year, items in answers.get(lecture_number, {}).items()
                    if str(int(match.group(1))) in items
                ]
                if len(candidate_years) == 1:
                    matched_year = candidate_years[0]
                    expected = answers[lecture_number][matched_year]
            if match and str(int(match.group(1))) in expected:
                finish_active()
                number = int(match.group(1))
                answer = expected[str(number)]
                first_entry = {
                    "type": "text",
                    "text": match.group(3).strip(),
                    "imageRels": list(block.get("imageRels", [])),
                }
                active = {
                    "lectureNumber": lecture_number,
                    "lectureTitle": lecture["title"],
                    "sourceLectureTitle": source_title,
                    "year": matched_year,
                    "number": number,
                    "answers": answer["answers"],
                    "answerStatus": answer["status"],
                    "answerRaw": answer["raw"],
                    "answerNote": answer["note"],
                    "professor2026": lecture["professors"].get("2026", ""),
                    "professorAtExam": lecture["professors"].get(year, ""),
                    "importance": importance(lecture, year),
                    "classificationStatus": "DOCX 강의 구획 기준",
                    "yearStatus": "정답표 기준 연도 교정" if matched_year != year else "DOCX 연도 표지 기준",
                    "pageHint": 3,
                    "entries": [first_entry],
                }
                continue
            if active:
                if block["type"] == "paragraph":
                    active["entries"].append(
                        {"type": "text", "text": block["text"], "imageRels": list(block.get("imageRels", []))}
                    )
                elif block["type"] == "table":
                    active["entries"].append(
                        {"type": "table", "rows": block["rows"], "imageRels": list(block.get("imageRels", []))}
                    )
        finish_active()

        found = {(question["year"], str(question["number"])) for question in questions if question["lectureNumber"] == lecture_number}
        for answer_year, answer_items in answers.get(lecture_number, {}).items():
            for answer_number in answer_items:
                if (answer_year, answer_number) not in found:
                    unresolved.append(
                        {
                            "lectureNumber": lecture_number,
                            "year": answer_year,
                            "number": int(answer_number),
                            "reason": "정답표에는 있으나 문제 시작 문단을 찾지 못함",
                        }
                    )
    return questions, unresolved


def main() -> None:
    document = Document(str(SOURCE_DOCX))
    blocks = make_blocks(document)
    years, lectures = professor_table(blocks)
    positions = header_positions(blocks)
    answer_start, answers = parse_answers(blocks, positions)
    questions, unresolved = parse_questions(document, blocks, lectures, positions, answer_start, answers)
    OUTPUT_DATA.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_REPORT.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "schemaVersion": 1,
        "source": SOURCE_DOCX.name,
        "years": years,
        "lectures": list(lectures.values()),
        "questions": questions,
    }
    OUTPUT_DATA.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    report = {
        "sourcePages": len(PdfReader(str(SOURCE_PDF)).pages),
        "documentParagraphs": len(document.paragraphs),
        "documentTables": len(document.tables),
        "inlineShapes": len(document.inline_shapes),
        "mediaRelationships": len({rel for question in questions for rel in question.get("assets", [])}),
        "lectureCount": len(lectures),
        "questionCount": len(questions),
        "answerCount": sum(len(items) for lecture in answers.values() for items in lecture.values()),
        "questionsWithFiveChoices": sum(len(question["choices"]) == 5 for question in questions),
        "questionsWithAssets": sum(bool(question["assets"]) for question in questions),
        "questionTableImages": sum(
            sum("-table-" in asset for asset in question.get("assets", []))
            for question in questions
        ),
        "unresolved": unresolved,
    }
    OUTPUT_REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
